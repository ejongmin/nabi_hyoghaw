# -*- coding: utf-8 -*-
"""
나비효과 캡스톤 최종 보고서 생성 스크립트
수자인팀 스타일을 따른 학술 보고서 포맷
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"C:/Users/john9/nabi_hyoghaw/reports/nabi_capstone_report_final.docx"

doc = Document()

# ---------- 기본 폰트 설정 ----------
KOREAN_FONT = "맑은 고딕"
BODY_SIZE = 11

def set_default_style():
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

def style_heading(level, size, bold=True):
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)

set_default_style()
style_heading(1, 18)
style_heading(2, 14)
style_heading(3, 12)

# ---------- 페이지 여백 설정 (A4, 2.5cm) ----------
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---------- 페이지 번호 (하단 중앙) ----------
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(10)

# ---------- 헬퍼 함수 ----------
def add_para(text, align=None, bold=False, size=None, space_after=6, line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), "Malgun Gothic")
    rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    return p

def add_heading_kr(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Malgun Gothic"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
        rfonts.set(qn("w:ascii"), "Malgun Gothic")
        rfonts.set(qn("w:hAnsi"), "Malgun Gothic")
    return h

def add_page_break():
    doc.add_page_break()

def add_table_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(11)
    run.font.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)

def add_table_source(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(9)
    run.font.italic = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)

def add_word_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(10)
        run.font.bold = True
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    # rows
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = "Malgun Gothic"
            run.font.size = Pt(10)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    return table

# ==========================================================
# 표지 페이지
# ==========================================================
# 상단 여백
for _ in range(4):
    add_para("", space_after=0)

add_para("DAT 6기 캡스톤 보고서", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=6)
add_para("", space_after=0)
add_para("", space_after=0)

# 한글 제목
add_para("EV 배터리 공급망 리스크의", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, space_after=4)
add_para("자본시장 파급 효과 실증 분석", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, space_after=10)
add_para("— 뉴스 감성과 그래프 어텐션 네트워크를 활용한", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=4)
add_para("공급망 전파 메커니즘 연구 —", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, space_after=12)

add_para("", space_after=0)

# 영문 부제
add_para("Empirical Analysis of Supply Chain Risk Spillover Effects",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
add_para("in the EV Battery Industry:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
add_para("A Study on Contagion Mechanisms Using News Sentiment",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
add_para("and Graph Attention Networks", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=10)

# 하단 여백
for _ in range(6):
    add_para("", space_after=0)

add_para("DAT 6기 나비효과(Nabi-Effect) 팀", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=8)
add_para("이 종 민", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=8)
add_para("2026년 5월", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=4)

add_page_break()

# ==========================================================
# 초록
# ==========================================================
add_heading_kr("초록 (Abstract)", level=1)

add_para(
    "본 연구는 전기차(EV) 배터리 공급망에서 발생하는 리스크 이벤트가 관련 기업의 주가에 미치는 "
    "파급 효과를 실증적으로 분석한다. 글로벌 뉴스 데이터베이스인 GDELT GKG로부터 2022~2025년 "
    "기사 4,894,621건을 수집하고, FinBERT 기반 감성 분석을 통해 40,928건의 리스크 이벤트를 "
    "추출하였다. 44개 EV 배터리 공급망 기업을 대상으로 누적 비정상 수익률(CAR) 이벤트 스터디를 "
    "수행한 결과, 직접 충격(1-hop)의 주가 반응은 통계적으로 유의하지 않았으나, 2-hop 간접 전파 "
    "효과는 강건하게 유의함을 확인하였다(MIXED_2: CAAR=+0.90%, t=3.70, p<0.001). 또한 공급망 "
    "지식 그래프 기반 그래프 어텐션 네트워크(GAT) Ablation 실험을 통해 실제 공급망 구조가 랜덤 "
    "그래프 및 순수 시계열 대비 우월한 예측 기여도를 나타냄을 확인하였다"
    "(AUC: 0.5281 vs. 0.5044 vs. 0.4594). 본 연구는 공급망 리스크의 네트워크 간접 전파 특성과 "
    "시장의 정보 마찰 가능성을 시사하며, EV 산업 투자자 및 정책 입안자에게 실질적 함의를 제공한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_para("", space_after=4)
add_para("Keywords: EV 배터리 공급망, 리스크 전파, CAR 이벤트 스터디, FinBERT, 그래프 어텐션 네트워크",
         bold=True, size=10)

add_para("", space_after=6)
add_heading_kr("Abstract", level=2)
add_para(
    "This study empirically analyzes how risk events in the EV battery supply chain propagate "
    "to the capital markets of related firms. Using 4,894,621 news articles from GDELT GKG "
    "(2022–2025) and FinBERT-based sentiment analysis, we identify 40,928 risk events and "
    "conduct a CAR event study on 44 global EV battery supply chain firms. Results show that "
    "while direct (1-hop) shock effects are statistically insignificant, 2-hop indirect "
    "contagion effects are robustly significant (MIXED_2: CAAR=+0.90%, t=3.70, p<0.001). "
    "An ablation study further confirms that the real supply chain graph structure outperforms "
    "random graphs and graph-free baselines in predictive performance "
    "(AUC: 0.5281 vs. 0.5044 vs. 0.4594). These findings suggest information frictions in "
    "supply chain risk propagation and provide actionable insights for EV industry investors "
    "and policymakers.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para("", space_after=4)
add_para("Keywords: EV battery supply chain, Risk contagion, CAR event study, FinBERT, "
         "Graph Attention Network", bold=True, size=10)

add_page_break()

# ==========================================================
# 목차
# ==========================================================
add_heading_kr("목   차", level=1)

toc_items = [
    ("초록", ""),
    ("제1장 서론", ""),
    ("    제1절 연구의 배경 및 목적", ""),
    ("    제2절 연구의 범위 및 방법", ""),
    ("제2장 이론적 배경", ""),
    ("    제1절 선행연구 검토", ""),
    ("    제2절 용어의 정의 및 개념적 틀", ""),
    ("제3장 연구 설계 및 분석 방법", ""),
    ("    제1절 연구 대상 및 자료 수집", ""),
    ("    제2절 분석 도구 및 절차", ""),
    ("제4장 연구 결과 및 실증 분석", ""),
    ("    제1절 자료의 특성", ""),
    ("    제2절 분석 결과", ""),
    ("    제3절 결과에 대한 논의", ""),
    ("제5장 결론", ""),
    ("    제1절 연구 요약", ""),
    ("    제2절 정책적 제언 및 시사점", ""),
    ("    제3절 연구의 한계 및 향후 과제", ""),
    ("참고문헌", ""),
]
for item, _ in toc_items:
    add_para(item, size=12, space_after=4, line_spacing=1.5)

add_page_break()

# 본문 시작 시점에 페이지 번호 적용을 위해 섹션 재설정
section = doc.sections[0]
add_page_number(section)

# ==========================================================
# 제1장 서론
# ==========================================================
add_heading_kr("제1장  서  론", level=1)

add_heading_kr("제1절  연구의 배경 및 목적", level=2)

add_heading_kr("1.1.1 연구 배경", level=3)
add_para(
    "전기차(Electric Vehicle, EV) 배터리 공급망은 리튬·코발트·니켈 등 핵심 광물의 채굴부터 "
    "양·음극재 가공, 배터리 셀 제조, 완성차 조립에 이르는 복잡한 다단계 글로벌 네트워크로 "
    "구성된다. 최근 급격히 확대된 EV 수요와 함께, 이 공급망은 지정학적 갈등, 환경 규제 강화, "
    "원자재 가격 변동 등 다양한 외부 충격에 노출되어 있으며, 공급망의 특정 지점에서 발생한 "
    "리스크가 전체 네트워크를 통해 전파되는 이른바 ‘나비효과(Butterfly Effect)’가 실질적 "
    "위협으로 부각되고 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "2022년 러시아-우크라이나 전쟁으로 인한 니켈 공급 충격, 2024~2025년 미-중 무역 갈등에 "
    "따른 희토류 수출 규제, 인도네시아의 니켈 원광 수출 금지 등 일련의 사건들은 공급망 단일 "
    "지점의 충격이 글로벌 OEM까지 연쇄적으로 영향을 미칠 수 있음을 실증하였다. 그러나 이러한 "
    "공급망 리스크의 자본시장 파급 경로와 전파 메커니즘에 대한 체계적 실증 연구는 상대적으로 "
    "부족한 실정이다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("1.1.2 연구 목적", level=3)
add_para(
    "본 연구는 다음의 세 가지 핵심 연구 질문을 제시한다. 첫째, EV 배터리 공급망에서 발생한 "
    "리스크 이벤트는 직접 영향 기업의 주가에 유의미한 비정상 수익률을 초래하는가? 둘째, 리스크 "
    "충격은 공급망 네트워크를 통해 2차·3차 연계 기업으로 전파되는가? 셋째, 공급망 그래프 구조 "
    "정보는 기업 주가의 이상 수익률 예측에 기여하는가? 이를 통해 EV 배터리 공급망 리스크의 "
    "체계적 파급 메커니즘을 규명하고, 투자자 및 정책 입안자를 위한 실질적 시사점을 도출하고자 "
    "한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("제2절  연구의 범위 및 방법", level=2)
add_para(
    "본 연구의 분석 대상은 전 세계 EV 배터리 공급망의 주요 상장 기업 44개사로, 업스트림"
    "(광산·정제, 2개사), 소재(양·음극재, 5개사), 셀/팩(배터리 제조, 5개사), OEM(완성차, 6개사) "
    "등 핵심 공급망 단계를 포괄하며, 분석 기간은 2017년 1월부터 2026년 5월까지이다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "분석 방법은 크게 세 단계로 구성된다. 첫째, GDELT GKG(Global Knowledge Graph) 뉴스 "
    "데이터베이스에서 공급망 관련 기업 기사를 수집하고 FinBERT 감성 분석 모델로 리스크 "
    "이벤트를 식별한다. 둘째, 시장 모델 기반 CAR(Cumulative Abnormal Return) 이벤트 스터디를 "
    "통해 직접·간접 주가 영향을 정량화한다. 셋째, 공급망 지식 그래프를 구축하고 Graph Attention "
    "Network(GAT)를 적용하여 공급망 구조 정보의 예측적 가치를 탐색적으로 검증한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

# ==========================================================
# 제2장 이론적 배경
# ==========================================================
add_heading_kr("제2장  이론적 배경", level=1)

add_heading_kr("제1절  선행연구 검토", level=2)

add_heading_kr("2.1.1 공급망 리스크와 자본시장 반응", level=3)
add_para(
    "공급망 리스크가 기업 가치에 미치는 영향 연구는 Hendricks and Singhal(2003)의 선구적 "
    "연구로부터 출발한다. 이들은 제조업 기업의 공급망 차질(supply chain disruption) 공시가 "
    "평균적으로 약 10%의 주가 하락을 초래함을 이벤트 스터디로 확인하였다. 이후 Bakshi and "
    "Kleindorfer(2009)는 공급망 취약성과 기업 가치 간의 이론적 관계를 규명하였으며, "
    "Jüttner(2005)는 공급망 리스크를 수요 리스크, 공급 리스크, 환경 리스크로 유형화하는 분류 "
    "체계를 제시하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "최근에는 뉴스 데이터를 활용한 자동화된 리스크 측정 연구가 확산되고 있다. Tetlock(2007)은 "
    "미디어 감성이 주가를 선행한다는 것을 실증하였으며, Baker et al.(2016)은 뉴스 기반 "
    "경제정책 불확실성(EPU) 지수를 개발하였다. GDELT 데이터를 활용한 연구로는 Leetaru and "
    "Schrodt(2013)의 기초 연구 이후 지정학적 리스크 분석이 다방면으로 이루어지고 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("2.1.2 공급망 네트워크와 전파 효과", level=3)
add_para(
    "공급망을 네트워크 관점에서 분석한 연구들은 리스크의 2차 전파(contagion) 가능성을 "
    "제시한다. Acemoglu et al.(2012)은 산업 간 연계 구조에서 국지적 충격이 전체 경제에 "
    "파급되는 메커니즘을 이론적으로 규명하였다. Carvalho(2014)는 투입-산출 네트워크 분석을 "
    "통해 허브 기업에서 발생한 충격이 하류 기업으로 폭포처럼 전파됨을 보였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("2.1.3 그래프 신경망의 금융 적용", level=3)
add_para(
    "그래프 신경망(GNN)의 금융 분야 적용은 최근 급증하고 있다. Kipf and Welling(2017)의 GCN을 "
    "시작으로, Veličković et al.(2018)의 GAT는 이웃 노드 간 가중치 차별화를 통해 금융 네트워크 "
    "분석에 광범위하게 활용된다. 본 연구는 선행연구를 EV 배터리라는 특정 공급망 도메인에 "
    "적용하고, 공급망 구조 자체의 기여도를 Ablation 실험으로 검증한다는 점에서 차별성을 "
    "지닌다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("제2절  용어의 정의 및 개념적 틀", level=2)

add_heading_kr("2.2.1 주요 용어 정의", level=3)
add_para(
    "누적 비정상 수익률(CAR, Cumulative Abnormal Return): 이벤트 창(event window) 기간 동안 "
    "특정 기업의 실제 수익률에서 시장 모델로 예측한 정상 수익률을 차감한 값의 합. 공급망 리스크 "
    "이벤트에 대한 시장의 가격 반응을 측정하는 핵심 지표이다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "공급망 리스크 이벤트(Supply Chain Risk Event): FinBERT 감성 분석 모델이 부정(NEG) 또는 "
    "긍정(POS) 감성 충격으로 분류한 뉴스 기사 클러스터. Z-score의 절댓값이 2.0을 초과하는 "
    "감성 시계열 극단값을 충격(shock)으로 식별한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "전파 관계(Contagion): 직접 충격을 받은 기업(DIRECT)으로부터 공급망 네트워크를 통해 "
    "연결된 기업들에게 CAR이 파급되는 현상. hop 수에 따라 1-hop(직접 연결), 2-hop(2단계 간접 "
    "연결)으로 구분한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "GDELT GKG(Global Knowledge Graph): 전 세계 뉴스 미디어를 실시간 수집·분석하는 오픈소스 "
    "데이터베이스로, 인물·조직·위치·테마 등의 개체 및 관계 정보를 구조화하여 제공한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("2.2.2 개념적 분석 틀", level=3)
add_para(
    "본 연구의 개념적 틀은 세 층위로 구성된다. 첫 번째 층위는 뉴스 데이터에서 리스크 이벤트를 "
    "추출하는 데이터 파이프라인, 두 번째 층위는 이벤트 스터디를 통한 자본시장 파급 효과 측정, "
    "세 번째 층위는 지식 그래프 기반 GNN을 통한 네트워크 전파 탐색이다. 이 세 층위는 공급망 "
    "리스크의 식별 → 측정 → 예측이라는 순차적 논리 구조를 형성한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

# ==========================================================
# 제3장 연구 설계 및 분석 방법
# ==========================================================
add_heading_kr("제3장  연구 설계 및 분석 방법", level=1)

add_heading_kr("제1절  연구 대상 및 자료 수집", level=2)

add_heading_kr("3.1.1 분석 기업 유니버스", level=3)
add_para(
    "본 연구의 분석 대상은 EV 배터리 공급망의 가치사슬 전반을 포괄하는 글로벌 상장 기업 44개사이다. "
    "기업 선정 기준은 (1) EV 배터리 가치사슬에서 매출의 30% 이상을 차지하는 사업 영역 보유, "
    "(2) 주요 거래소 상장, (3) 2017년 이후 연속 주가 데이터 이용 가능으로 한정하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_table_title("<표 1> EV 배터리 공급망 핵심 기업 유니버스(18개사)")
add_word_table(
    ["공급망 단계", "기업명", "티커", "상장거래소"],
    [
        ["Upstream", "Albemarle", "ALB", "NYSE"],
        ["Upstream", "Glencore", "GLEN.L", "LSE"],
        ["Materials", "BASF", "BAS.DE", "FSE"],
        ["Materials", "CNGR Advanced Material", "300919.SZ", "SZSE"],
        ["Materials", "EcoPro BM", "247540.KQ", "KRX"],
        ["Materials", "Huayou Cobalt", "603799.SH", "SSE"],
        ["Materials", "LG Chem", "051910.KS", "KRX"],
        ["Cell/Pack", "CALB", "3931.HK", "SEHK"],
        ["Cell/Pack", "CATL", "300750.SZ", "SZSE"],
        ["Cell/Pack", "EVE Energy", "300014.SZ", "SZSE"],
        ["Cell/Pack", "Gotion High-Tech", "002074.SZ", "SZSE"],
        ["Cell/Pack", "LG Energy Solution", "373220.KS", "KRX"],
        ["OEM", "BMW", "BMW.DE", "FSE"],
        ["OEM", "BYD", "1211.HK", "SEHK"],
        ["OEM", "Honda", "7267.T", "TSE"],
        ["OEM", "Mercedes-Benz", "MBG.DE", "FSE"],
        ["OEM", "Tesla", "TSLA", "NASDAQ"],
        ["OEM", "Volkswagen", "VOW3.DE", "FSE"],
    ]
)
add_table_source("출처: 저자 작성 (기업 IR, 연간보고서, 공급망 데이터베이스 종합)")

add_heading_kr("3.1.2 뉴스 데이터 수집", level=3)
add_para(
    "뉴스 데이터는 GDELT GKG 데이터베이스에서 수집하였다. 수집 기간은 2022년 1월부터 2025년 "
    "12월까지이며, BigQuery API를 활용하여 분석 대상 기업명 및 관련 키워드가 포함된 기사를 "
    "추출하였다. 원본 수집 기사 수는 총 4,894,621건이며, 중복 제거·관련성 필터링·품질 검증 "
    "과정을 거쳐 최종 40,928건의 리스크 이벤트를 도출하였다(압축률 99.2%).",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("3.1.3 주가 데이터", level=3)
add_para(
    "주가 데이터는 Yahoo Finance API를 통해 수집한 일별 수정 종가를 사용하였다. 수집 기간은 "
    "2017년 1월부터 2026년 5월까지이며, 각 기업의 거래소별 시장 인덱스(한국=KOSPI, "
    "중국=CSI300, 일본=Nikkei225, 유럽=DAX, 미국=S&P500, 홍콩=HSI)를 시장 수익률 추정에 "
    "활용하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("제2절  분석 도구 및 절차", level=2)

add_heading_kr("3.2.1 FinBERT 감성 분석 및 이벤트 식별", level=3)
add_para(
    "수집된 뉴스 기사의 감성 분석에는 FinBERT(Araci, 2019)를 사용하였다. FinBERT는 BERT "
    "사전학습 언어 모델을 금융 도메인 텍스트로 파인튜닝한 모델로, 금융 뉴스에서 긍정/부정/중립 "
    "감성을 높은 정확도로 분류한다. 영어 이외의 언어로 작성된 기사에는 각 언어에 특화된 다국어 "
    "FinBERT 모델을 적용하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "이벤트 식별 절차는 다음과 같다. ① 기업별 월별 감성 점수의 가중 평균 산출, ② 36개월 롤링 "
    "평균과 표준편차 기반 Z-score 정규화, ③ 절댓값 Z-score ≥ 2.0인 월을 감성 충격 이벤트로 "
    "정의 (NEG: Z < -2.0, POS: Z > +2.0). 최종적으로 18개 기업 기준 101개월 시계열에서 NEG "
    "충격 37건, POS 충격 28건이 식별되었다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("3.2.2 CAR 이벤트 스터디", level=3)
add_para(
    "이벤트 스터디는 Brown and Warner(1985)의 시장 모델을 기반으로 수행하였다. 정상 수익률 "
    "추정을 위한 추정 창(estimation window)은 이벤트 발생 전 [-120, -21] 영업일로 설정하였으며, "
    "최소 60거래일 이상의 데이터가 있는 경우에만 유효한 이벤트로 인정하였다. 주 분석 이벤트 "
    "창은 [0, +21] 영업일(약 1개월)로 설정하였으며, 추가적으로 [-1,+1], [0,+1], [0,+5], "
    "[0,+10], [-5,+5] 창을 보조 분석에 활용하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "통계 검정은 횡단면 t-검정(cross-sectional t-test)과 부호 검정(sign test)을 병행하였으며, "
    "표준 오차의 견고성을 위해 이벤트-월 클러스터 부트스트랩(B=1,000)을 적용하였다. 공급망 "
    "전파 분석을 위해 직접 충격 기업(DIRECT)과 공급망 연결 기업의 CAR을 hop 수에 따라 "
    "구분하여 분석하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("3.2.3 공급망 지식 그래프 구축", level=3)
add_para(
    "공급망 지식 그래프는 학술 문헌, 기업 연간보고서, 공급망 데이터베이스를 수동 검토하여 "
    "구축하였다. 18개 핵심 기업을 노드로, 65개 엣지(SUPPLIES 50건, PARTNERS_WITH 7건, "
    "COMPETES_WITH 5건, 기타 3건)를 관계로 구성하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("3.2.4 그래프 어텐션 네트워크(GAT) 실험", level=3)
add_para(
    "공급망 구조 정보의 예측 기여도 검증을 위해 Veličković et al.(2018)의 GAT 모델을 "
    "적용하였다. 입력 피처는 FinBERT 감성 Z-score를 포함한 18개 변수이며, 예측 목표는 다음 월의 "
    "비정상 수익률 방향(상승/하락)이다. 실험 설계는 훈련(2017-01~2023-06), "
    "검증(2023-07~2024-06), 테스트(2024-07~2026-05)의 워크포워드 교차 검증을 적용하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "Ablation Study 설계: ① Real Graph GAT(실제 공급망 구조), ② Random Graph GAT(무작위 그래프, "
    "구조 기여도 검증용), ③ No Graph GRU(그래프 미사용 순수 시계열 베이스라인)를 비교하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

# ==========================================================
# 제4장 연구 결과 및 실증 분석
# ==========================================================
add_heading_kr("제4장  연구 결과 및 실증 분석", level=1)

add_heading_kr("제1절  자료의 특성", level=2)

add_heading_kr("4.1.1 뉴스 이벤트 분포", level=3)

add_table_title("<표 2> 이벤트 식별 단계별 요약")
add_word_table(
    ["구분", "건수"],
    [
        ["GDELT 원본 수집 기사", "4,894,621건"],
        ["기업 관련 필터링 후 리스크 이벤트", "40,928건"],
        ["FinBERT 이벤트 식별 (18개 기업)", "65건 (NEG: 37, POS: 28)"],
        ["CAR 분석 유효 이벤트 (44개 기업)", "129건 (NEG: 75, POS: 61)"],
    ]
)
add_table_source("출처: 저자 작성 (GDELT GKG, FinBERT 분석 결과)")

add_heading_kr("4.1.2 주가 데이터 특성", level=3)
add_para(
    "44개 기업의 일별 주가 데이터는 총 92,350건의 수익률 관측치를 포함하며, 결측률은 대부분의 "
    "기업에서 0%로 양호하다. 수익률의 절댓값이 20%를 초과하는 극단값 비율은 전체의 0.1% "
    "미만으로, 데이터 품질 기준을 충족한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("제2절  분석 결과", level=2)

add_heading_kr("4.2.1 직접 효과(Direct Effect) — CAR 이벤트 스터디", level=3)

add_table_title("<표 3> 이벤트 창별 CAAR 추정 결과 (DIRECT, N=129)")
add_word_table(
    ["이벤트 창", "N", "CAAR", "t-통계량", "p값"],
    [
        ["[-1, +1]", "129", "+0.29%", "0.637", "0.525"],
        ["[0, +1]",  "129", "+0.11%", "0.269", "0.789"],
        ["[0, +5]",  "129", "+0.15%", "0.236", "0.814"],
        ["[0, +10]", "129", "-0.66%", "-0.820", "0.414"],
        ["[0, +21]", "125", "+0.62%", "0.608", "0.544"],
        ["[-5, +5]", "129", "+0.09%", "0.122", "0.903"],
    ]
)
add_table_source("출처: 저자 산출 (Phase 4 CAR Event Study, 시장 모델 추정 [-120,-21])")

add_para("충격 유형별 이질성 ([0, +21] 창):", bold=True)
add_para("· NEG shock: CAAR=-1.73%, t=-1.397, p=0.167 (비유의)")
add_para("· POS shock: CAAR=+3.62%, t=2.227, p=0.030** (유의)")
add_para(
    "주 분석 창인 [0,+21] 영업일에서 전체 이벤트 기준 CAAR=+0.62%(p=0.544)로 통계적으로 "
    "유의하지 않았다. 충격 유형별로는 긍정 충격(POS)의 경우 CAAR=+3.62%(p=0.030)로만 유의성이 "
    "확인되었다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("4.2.2 공급망 전파 효과(Contagion Effect)", level=3)

add_table_title("<표 4> 공급망 hop 거리별 전파 효과 (이벤트 창 [0, +21])")
add_word_table(
    ["전파 유형", "hop", "N", "CAAR", "t-통계량", "p값", "유의성"],
    [
        ["DIRECT", "-", "125", "+0.62%", "0.608", "0.544", "-"],
        ["DOWNSTREAM", "1", "161", "-0.43%", "-0.442", "0.659", "-"],
        ["UPSTREAM", "1", "203", "+0.05%", "0.060", "0.952", "-"],
        ["PEER", "1", "96", "+0.58%", "0.391", "0.697", "-"],
        ["DOWNSTREAM_2", "2", "384", "-0.49%", "-0.764", "0.445", "-"],
        ["MIXED_2", "2", "3,356", "+0.90%", "3.698", "<0.001", "***"],
        ["PEER_2", "2", "44", "+5.11%", "1.803", "0.078", "*"],
        ["UPSTREAM_2", "2", "898", "+1.17%", "2.300", "0.022", "**"],
    ]
)
add_table_source("출처: 저자 산출 (Phase 4 Contagion Analysis, *** p<0.001, ** p<0.05, * p<0.10)")

add_heading_kr("4.2.3 강건성 검정(Robustness Check)", level=3)
add_para(
    "T0 대안 설계(T0=M) 적용 결과, NEG shock에서 CAAR=-3.18%(p=0.070*)로 방향이 역전되며 약한 "
    "유의성이 확인되었다. 클러스터 표준오차(이벤트-월 클러스터, B=1,000) 적용 후에는 p=0.124로 "
    "유의성이 약화되어 결과의 통계적 신뢰성 측면의 한계가 존재한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("4.2.4 GAT Ablation Study", level=3)

add_table_title("<표 5> GAT Ablation Study 성능 비교 (테스트 2024-07~2026-05)")
add_word_table(
    ["모델", "AUC", "F1", "Acc", "IC월평균", "IC>0비율"],
    [
        ["Real Graph GAT", "0.5281", "0.5667", "48.0%", "+0.166", "67%"],
        ["Random Graph GAT", "0.5044", "0.5778", "46.8%", "-0.104", "20%"],
        ["No Graph GRU", "0.4594", "0.4825", "46.8%", "-0.002", "60%"],
        ["Dynamic GAT", "0.4495", "0.4188", "45.6%", "-0.133", "47%"],
    ]
)
add_table_source(
    "출처: 저자 산출 (Phase 5 Ablation, 18개 기업, 23개월, Permutation test IC p=0.398)"
)

add_para(
    "GAT 어텐션 가중치 상위 기여 엣지: CNGR↔알베마를(0.149), BYD↔혼다(0.148), "
    "EcoPro BM↔화유코발트(0.141)."
)

add_heading_kr("제3절  결과에 대한 논의", level=2)
add_para(
    "첫째, 1-hop 직접 효과 부재 및 2-hop 간접 전파의 강건성은 시장의 정보 마찰(information "
    "friction) 가능성을 시사한다. EV 배터리 공급망의 복잡한 연결 구조에서, 직접 충격을 받은 "
    "기업의 1차 거래 파트너에 대한 시장의 가격 반응은 즉각적으로 형성되지 않는 반면, 2단계 이상 "
    "거리에 위치한 기업에 대해서는 통계적으로 유의한 반응이 나타났다. 이는 Carvalho(2014)의 "
    "투입-산출 네트워크 이론에서 예측한 간접 전파 메커니즘과 부합하며, 투자자들이 공급망 2차 "
    "연쇄 효과를 지연 반영하는 행태적 편향의 증거일 수 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "둘째, 긍정 충격(POS)의 유의성은 EV 배터리 공급망 기업에 대한 호재 정보의 자본시장 즉각 "
    "반영이 부정 충격보다 상대적으로 효율적임을 시사한다. EV 산업의 성장 기대감이 긍정 정보의 "
    "신속한 가격 반영을 촉진하는 구조적 요인으로 작용할 가능성이 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "셋째, 그래프 구조의 탐색적 기여도는 공급망 네트워크 정보가 개별 기업의 시계열 정보를 "
    "보완하는 예측 신호를 제공함을 시사한다. 그러나 Permutation test p=0.398로 통계적 유의성은 "
    "확인되지 않아, 소규모 유니버스(18개 기업)에서는 탐색적 수준의 결론으로 해석해야 한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

# ==========================================================
# 제5장 결론
# ==========================================================
add_heading_kr("제5장  결  론", level=1)

add_heading_kr("제1절  연구 요약", level=2)
add_para(
    "본 연구는 EV 배터리 공급망 리스크의 자본시장 파급 메커니즘을 실증적으로 규명하기 위해 "
    "GDELT 뉴스 기반 리스크 이벤트 식별, CAR 이벤트 스터디, 그래프 어텐션 네트워크를 활용한 "
    "End-to-End 분석 파이프라인을 구축하였다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para("핵심 결과는 다음과 같다.", bold=True)
add_para(
    "① 긍정 충격에 대한 직접 반응(CAAR=+3.62%, p=0.030)은 유의하였으나, 부정 충격 직접 효과는 "
    "비유의하였다."
)
add_para(
    "② 2-hop 간접 전파에서 MIXED_2(CAAR=+0.90%, p<0.001)와 UPSTREAM_2(CAAR=+1.17%, p=0.022)의 "
    "강건한 유의성이 확인되었다."
)
add_para(
    "③ Ablation Study를 통해 실제 공급망 그래프 구조(AUC=0.5281)가 랜덤 그래프(0.5044) 및 "
    "그래프 미사용 모델(0.4594)보다 우월한 예측 성능을 보였다."
)

add_heading_kr("제2절  정책적 제언 및 시사점", level=2)
add_para(
    "투자자 관점: EV 배터리 공급망 리스크는 직접 충격 기업뿐 아니라 2-hop 간접 연결 기업의 "
    "포트폴리오에도 유의미한 영향을 미친다. 투자자들은 EV 관련 기업 분석 시 직접 거래 파트너를 "
    "넘어 2단계 이상의 공급망 연결성을 고려한 리스크 모니터링 체계를 갖출 필요가 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "기업 관점: 공급망 가시성(supply chain visibility) 확보가 리스크 관리의 핵심 과제임을 "
    "시사한다. UPSTREAM_2 기업군에서 양(+)의 비정상 수익률이 관찰된 결과는, 직접 충격 기업의 "
    "대안 공급원으로서의 역할이 자본시장에서 인식됨을 의미한다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "정책 관점: EV 배터리 핵심 광물의 특정 국가·기업 집중도가 높은 현황에서, 공급망 다변화 "
    "정책은 시스템적 리스크 전파를 억제하는 효과를 지닐 수 있다.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)

add_heading_kr("제3절  연구의 한계 및 향후 과제", level=2)
add_para("연구의 한계는 다음과 같다.", bold=True)
add_para(
    "① 44개 기업·129건 이벤트의 표본 규모 한계 — 클러스터 SE 적용 후 유의성 약화.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "② 월별 감성 충격 기반 이벤트 정의의 날짜 특정 모호성.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "③ 18개 기업 소규모 네트워크에서 GNN의 통계적 유의성 확보 한계.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    "④ 상관관계 기반 분석으로 인과 관계 규명의 한계.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para("향후 과제는 다음과 같다.", bold=True)
add_para("① 일별 뉴스-주가 고빈도 분석을 통한 식별력 강화,",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para("② 100개 이상 기업으로의 공급망 유니버스 확장,",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para("③ Temporal GNN(TGAT, DyRep) 적용을 통한 동적 전파 모델링,",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para("④ 반도체·희토류 등 타 산업으로의 방법론 일반화.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ==========================================================
# 참고문헌
# ==========================================================
add_heading_kr("참고문헌", level=1)

refs = [
    "Acemoglu, D., Carvalho, V. M., Ozdaglar, A., & Tahbaz-Salehi, A. (2012). The network "
    "origins of aggregate fluctuations. Econometrica, 80(5), 1977–2016.",
    "Araci, D. (2019). FinBERT: Financial sentiment analysis with pre-trained language "
    "models. arXiv preprint arXiv:1908.10063.",
    "Baker, S. R., Bloom, N., & Davis, S. J. (2016). Measuring economic policy uncertainty. "
    "The Quarterly Journal of Economics, 131(4), 1593–1636.",
    "Bakshi, N., & Kleindorfer, P. (2009). Co-opetition and investment for supply-chain "
    "resilience. Production and Operations Management, 18(6), 583–603.",
    "Brown, S. J., & Warner, J. B. (1985). Using daily stock returns: The case of event "
    "studies. Journal of Financial Economics, 14(1), 3–31.",
    "Carvalho, V. M. (2014). From micro to macro via production networks. Journal of "
    "Economic Perspectives, 28(4), 23–48.",
    "Gulley, A. L., Nassar, N. T., & Phillip, S. (2018). China, the United States, and "
    "competition for resources that enable emerging technologies. Proceedings of the "
    "National Academy of Sciences, 115(16), 4111–4115.",
    "Hendricks, K. B., & Singhal, V. R. (2003). The effect of supply chain glitches on "
    "shareholder wealth. Journal of Operations Management, 21(5), 501–522.",
    "Jüttner, U. (2005). Supply chain risk management: Understanding the business "
    "requirements from a practitioner perspective. The International Journal of Logistics "
    "Management, 16(1), 120–141.",
    "Kipf, T. N., & Welling, M. (2017). Semi-supervised classification with graph "
    "convolutional networks. ICLR 2017.",
    "Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global data on events, location, and "
    "tone, 1979–2012. ISA Annual Convention.",
    "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in "
    "the stock market. The Journal of Finance, 62(3), 1139–1168.",
    "Veličković, P., Cucurull, G., Casanova, A., Romero, A., Lio, P., & Bengio, Y. (2018). "
    "Graph attention networks. ICLR 2018.",
]
for r in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(-0.8)
    run = p.add_run(r)
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(11)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)

# ---------- 저장 ----------
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)

size = os.path.getsize(OUTPUT)
print(f"[OK] saved: {OUTPUT}")
print(f"[OK] size : {size:,} bytes ({size/1024:.1f} KB)")
