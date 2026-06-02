"""
나비효과 캡스톤 보고서 - 양식 적용 빌더
- 템플릿: DAT_캡스톤_프로젝트_보고서_양식.docx
- 내용: nabi_capstone_report_final.docx 전문 + figures/
- 출력: nabi_capstone_report_template_v1.docx
"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = Path("C:/Users/john9/Desktop/dat/DAT_캡스톤_프로젝트_보고서_양식.docx")
FIGURES  = Path("C:/Users/john9/nabi_hyoghaw/reports/figures")
OUTPUT   = Path("C:/Users/john9/nabi_hyoghaw/reports/nabi_capstone_report_template_v2.docx")

# ── 폰트 상수 ──────────────────────────────────────────────────────────────────
F_BOLD   = "한국외대체 B"
F_MED    = "한국외대체 M"
F_LIGHT  = "한국외대체 L"

# ── 헬퍼 ───────────────────────────────────────────────────────────────────────
def _set_run(run, text, font_name, size_pt, bold=False, italic=False, color=None):
    run.text = text
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, font_name=F_LIGHT, size_pt=11,
             bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=4, line_spacing=Pt(18)):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        _set_run(r, text, font_name, size_pt, bold=bold)
    return p

def add_heading1(doc, text):
    """제 X장  …"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    _set_run(r, text, F_BOLD, 16)

def add_heading2(doc, text):
    """제 X절  …"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    _set_run(r, text, F_MED, 14)

def add_heading3(doc, text):
    """X.X.X  …"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    _set_run(r, text, F_MED, 12, bold=True)

def add_body(doc, text):
    add_para(doc, text, F_LIGHT, 11)

def add_caption(doc, text):
    p = add_para(doc, text, F_MED, 10, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=2, space_after=2)
    return p

def add_figure(doc, filename, caption_text, desc_text=None, width_cm=14):
    """그림 삽입 + 캡션 + 선택적 부가 설명"""
    fig_path = FIGURES / filename
    if not fig_path.exists():
        add_body(doc, f"[그림 누락: {filename}]")
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(str(fig_path), width=Cm(width_cm))
    add_caption(doc, caption_text)
    if desc_text:
        pd = add_para(doc, desc_text, F_LIGHT, 10,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_before=1, space_after=6, line_spacing=Pt(16))
        pd.paragraph_format.left_indent  = Cm(0.5)
        pd.paragraph_format.right_indent = Cm(0.5)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type("page"))

def docx_break_type(btype="page"):
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE

def insert_page_break(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbk = OxmlElement('w:pageBreakBefore')
    pbk.set(qn('w:val'), '1')
    pPr.append(pbk)

def add_table_from_data(doc, headers, rows, caption=""):
    """헤더 + 데이터행으로 표 생성"""
    if caption:
        add_caption(doc, caption)
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    # 헤더
    hdr = tbl.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.name = F_MED
            run.font.size = Pt(10)
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 데이터
    for i, row_data in enumerate(rows):
        row = tbl.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = F_LIGHT
                run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 여백
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  메인 빌더
# ══════════════════════════════════════════════════════════════════════════════
def build():
    shutil.copy(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)

    # ── 1. 표지 플레이스홀더 교체 ─────────────────────────────────────────────
    replacements = {
        "DAT N기 캡스톤 프로젝트 보고서": "DAT 7기 캡스톤 프로젝트 보고서",
        "제목": "EV 배터리 공급망 리스크의 자본시장 파급 효과 실증 분석",
        ": 부제목": "— 뉴스 감성과 그래프 어텐션 네트워크를 활용한 공급망 전파 메커니즘 연구 —",
        "영문제목": "Empirical Analysis of Supply Chain Risk Spillover in the EV Battery Industry",
        "보고서 요약": (
            "본 연구는 전기차(EV) 배터리 공급망에서 발생하는 리스크 이벤트가 관련 기업 주가에 미치는 파급 효과를 실증적으로 분석한다. "
            "GDELT GKG에서 3,841,469건의 뉴스 기사에 FinBERT 감성 분석을 적용하여 리스크 이벤트를 식별하고, "
            "44개 글로벌 기업(69 노드·287 엣지 공급망 지식 그래프)을 대상으로 CAR 이벤트 스터디로 전파 메커니즘을 검증하였다. "
            "핵심 결과로, 긍정 충격에 대한 직접 반응(CAAR=+3.62%, p=0.030)은 유의하였으나 부정 충격 직접 효과는 비유의하였다. "
            "2-hop 간접 전파에서 MIXED_2(CAAR=+0.90%, p<0.001)와 UPSTREAM_2(CAAR=+1.17%, p=0.022)의 "
            "강건한 유의성이 확인되어 정보 마찰(information friction) 가설을 지지한다. "
            "Ablation Study를 통해 실제 공급망 그래프 구조를 활용한 GAT(AUC=0.5281)가 랜덤 그래프(0.5044)보다 우월한 예측 성능을 보였다."
        ),
        "보고서 요약(영문)": (
            "This study empirically analyzes how risk events in the EV battery supply chain propagate to capital markets. "
            "We applied FinBERT sentiment analysis to 3,841,469 GDELT GKG articles, identified sentiment shocks, and quantified "
            "contagion effects via CAR event studies across 44 global firms (knowledge graph: 69 nodes, 287 edges). "
            "Positive shocks yielded significant direct responses (CAAR=+3.62%, p=0.030), while negative direct effects were insignificant. "
            "Robust 2-hop indirect contagion was observed in MIXED_2 (CAAR=+0.90%, p<0.001) and "
            "UPSTREAM_2 (CAAR=+1.17%, p=0.022), supporting the information-friction hypothesis. "
            "A Graph Attention Network (GAT) ablation confirmed that real supply chain graph structure (AUC=0.5281) "
            "outperforms random graphs (0.5044), suggesting exploratory predictive value of network topology."
        ),
        "DAT N기": "DAT 7기",
        "팀 명": "나비효과(Nabi-Effect) 팀",
        "팀 구성원": "이 종 민",
        "2023년 06월 08일": "2026년 5월",
    }

    for para in doc.paragraphs:
        full = para.text.strip()
        if full in replacements:
            new_text = replacements[full]
            # 모든 run 지우고 첫 run만 남김
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    # ── 2. 목차 교체 (p67~97 구간: 기존 목차 텍스트를 실제 목차로 교체) ────────
    toc_texts = {
        "목 차(예시)": "목   차",
        "제 1장 서론 \t1": "제1장  서  론\t1",
        "제 1절 연구 배경 \t1": "    제1절  연구의 배경 및 목적\t1",
        "제 2절 연구 목적 및 논문 구성 \t2": "    제2절  연구의 범위 및 방법\t2",
        "제 2장 데이터 소개\t2": "제2장  이론적 배경\t3",
        "제 1절 선행연구 분석\t2": "    제1절  선행연구 검토\t3",
        "제 2절 조사 및 분석 데이터 소개\t3": "    제2절  용어의 정의 및 개념적 틀\t5",
        "제 3장 분석 결과\t6": "제3장  연구 설계 및 분석 방법\t6",
        "제 1절 연구 모형 설계\t6": "    제1절  연구 대상 및 자료 수집\t6",
        "제 2절 회귀계수 분석\t7": "    제2절  분석 도구 및 절차\t8",
        "제 4장 결론\t8": "제4장  연구 결과 및 실증 분석\t11",
        "제 1절 연구 결과 종합 및 시사점 도출\t8": "    제1절  자료의 특성\t11",
        "제 2절 한계점 및 향후 연구 방향성 제언\t9": "    제2절  분석 결과\t12",
        "참고문헌\t11": "    제3절  결과에 대한 논의\t16\n제5장  결  론\t18",
        "국내문헌\t11": "    제1절  연구 요약\t18",
        "국외문헌\t12": "    제2절  정책적 제언 및 시사점\t19",
        "표 목차": "    제3절  연구의 한계 및 향후 과제\t20\n참고문헌\t21",
        "[표 2-1]\t4": "[표 1] EV 배터리 공급망 핵심 기업 유니버스(18개사)\t7",
        "[표 2-2]\t5": "[표 2] 이벤트 식별 단계별 요약\t11",
        "[표 2-3]\t6": "[표 3] 이벤트 창별 CAAR 추정 결과 (DIRECT)\t12",
        "[표 3-1]\t7": "[표 4] 공급망 hop 거리별 전파 효과\t14",
        "그림 목차": "[표 5] GAT Ablation Study 성능 비교\t15\n\n그림 목차",
        "[그림 2-1]\t5": "[그림 1] EV 배터리 공급망 지식 그래프\t10",
        "[그림 3-1]\t6": "[그림 2] CAR 이벤트 창별 CAAR 추이\t13\n[그림 3] 공급망 hop별 전파 효과\t14\n[그림 4] GAT Ablation 결과 비교\t15\n[그림 5] GAT 어텐션 가중치 분포\t16",
    }

    for para in doc.paragraphs:
        full = para.text.strip()
        if full in toc_texts:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = toc_texts[full]
            else:
                para.add_run(toc_texts[full])

    # ── 3. 본문 영역 초기화 및 내용 삽입 ────────────────────────────────────────
    # 본문은 p98~ (제 1장 서론, 제 1절 연구 배경, 참고문헌 예시)
    # 해당 단락들을 삭제하고 실제 내용으로 교체
    body_elem = doc.element.body
    paras = doc.paragraphs
    # p98부터 마지막 sectPr 이전까지 삭제
    start_idx = 98
    para_elems = [p._p for p in paras[start_idx:]]
    for pe in para_elems:
        body_elem.remove(pe)

    # 이제 doc에 본문 단락을 추가
    _write_body(doc)

    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")


def _write_body(doc):
    """실제 본문 작성 (v2: 데이터셋 상세, 베이스라인 비교, 그림 부가설명 추가)"""

    # ══════════════════════════════════════════════════════════════════════════
    # 제1장 서론
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "제1장  서  론")
    add_heading2(doc, "제1절  연구의 배경 및 목적")
    add_heading3(doc, "1.1.1 연구 배경")
    add_body(doc,
        "전기차(Electric Vehicle, EV) 배터리 공급망은 리튬·코발트·니켈 등 핵심 광물의 채굴부터 양·음극재 가공, "
        "배터리 셀 제조, 완성차 조립에 이르는 복잡한 다단계 글로벌 네트워크로 구성된다. "
        "최근 급격한 EV 수요 성장과 함께 공급망의 지정학적 리스크 노출도가 크게 높아졌다."
    )
    add_body(doc,
        "2022년 러시아-우크라이나 전쟁으로 인한 니켈 공급 충격, 2024~2025년 미-중 무역 갈등에 따른 희토류 수출 규제, "
        "인도네시아의 니켈 원광 수출 금지 등 일련의 사건들은 공급망 단일 지점의 충격이 글로벌 OEM·소재 기업의 "
        "주가에까지 파급되는 리스크 전파(contagion) 현상을 야기하였다. 그러나 이러한 전파 메커니즘을 "
        "뉴스 기반 빅데이터와 그래프 모델을 결합하여 체계적으로 실증한 연구는 아직 드물다."
    )

    add_heading3(doc, "1.1.2 연구 목적")
    add_body(doc,
        "본 연구는 다음의 세 가지 핵심 연구 질문을 제시한다. 첫째, EV 배터리 공급망에서 발생한 리스크 이벤트는 "
        "직접 영향 기업의 주가에 유의미한 비정상 수익률을 초래하는가? 둘째, 리스크 충격은 공급망 네트워크를 통해 "
        "2차·3차 연결 기업으로 전파되는가? 셋째, 공급망 그래프 구조 정보는 기업 주가 수익률 예측에 기여하는가?"
    )
    add_body(doc,
        "이를 위해 본 연구는 GDELT 뉴스 기반 리스크 이벤트 식별 → CAR 이벤트 스터디 → "
        "그래프 어텐션 네트워크(GAT) 예측으로 이어지는 End-to-End 분석 파이프라인을 구축하고, "
        "공급망 리스크 전파의 실증적 근거를 제시한다."
    )

    add_heading2(doc, "제2절  연구의 범위 및 방법")
    add_body(doc,
        "본 연구의 분석 대상은 전 세계 EV 배터리 공급망의 주요 상장 기업 44개사로, "
        "업스트림(광산·정제, 2개사), 소재(양·음극재, 5개사), 셀/팩(배터리 제조, 5개사), OEM(완성차, 6개사) 등 "
        "핵심 공급망 참여자를 망라한다. 분석 기간은 뉴스 데이터 기준 2022~2025년, "
        "주가 데이터 기준 2017~2026년 5월이다."
    )
    add_body(doc,
        "분석 방법은 크게 세 단계로 구성된다. 첫째, GDELT GKG 뉴스 데이터베이스에서 공급망 관련 기업 기사를 수집하고 "
        "FinBERT 감성 분석 모델로 리스크 이벤트를 식별한다. 둘째, 시장 모델 기반 CAR 이벤트 스터디로 직접 효과와 "
        "hop 거리별 전파 효과를 추정한다. 셋째, 공급망 지식 그래프를 구축하고 GAT Ablation Study를 통해 "
        "그래프 구조 정보의 예측 기여도를 검증한다."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 제2장 이론적 배경
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "제2장  이론적 배경")
    add_heading2(doc, "제1절  선행연구 검토")
    add_heading3(doc, "2.1.1 공급망 리스크와 자본시장 반응")
    add_body(doc,
        "공급망 리스크가 기업 가치에 미치는 영향 연구는 Hendricks and Singhal(2003)의 선구적 연구로부터 출발한다. "
        "이들은 제조업 기업의 공급망 차질(supply chain disruption) 공시가 주가 하락과 유의미하게 연관됨을 발견하였다. "
        "이후 Carvalho(2014)는 생산 네트워크의 관점에서 국지적 충격이 상류·하류 기업으로 전파되는 메커니즘을 이론화하였다."
    )
    add_body(doc,
        "최근에는 뉴스 데이터를 활용한 자동화된 리스크 측정 연구가 확산되고 있다. Tetlock(2007)은 미디어 감성이 주가를 "
        "선행한다는 것을 실증하였으며, Baker et al.(2016)은 뉴스 기반 경제정책 불확실성 지수(EPU)가 투자·소비에 미치는 "
        "영향을 분석하였다. EV 배터리 분야에서는 Gulley et al.(2018)이 핵심 광물 공급 위험과 자본시장 연계성을 분석하였다."
    )

    add_heading3(doc, "2.1.2 공급망 네트워크와 전파 효과")
    add_body(doc,
        "공급망을 네트워크 관점에서 분석한 연구들은 리스크의 2차 전파(contagion) 가능성을 제시한다. "
        "Acemoglu et al.(2012)은 산업 간 연계 구조에서 국지적 충격이 전체 경제에 파급되는 메커니즘을 이론화하였다. "
        "Bakshi and Kleindorfer(2009)는 공급망 파트너 간 리스크 분담(co-opetition) 구조가 회복탄력성에 미치는 영향을 "
        "분석하였으며, Jüttner(2005)는 공급망 리스크 관리의 실무적 요건을 체계화하였다."
    )

    add_heading3(doc, "2.1.3 그래프 신경망의 금융 적용")
    add_body(doc,
        "그래프 신경망(GNN)의 금융 분야 적용은 최근 급증하고 있다. Kipf and Welling(2017)의 GCN을 시작으로, "
        "Veličković et al.(2018)의 GAT는 이웃 노드 간 가중치 차별화를 통해 금융 네트워크의 이질적 연결 구조를 "
        "포착하는 데 적합함을 보였다. 기업 네트워크에 GNN을 적용한 연구들은 주가 공동 움직임(co-movement)과 공급망 "
        "연결성 사이의 관계를 탐색하고 있으나, EV 배터리 공급망에 특화된 실증 연구는 아직 드물다."
    )

    add_heading2(doc, "제2절  용어의 정의 및 개념적 틀")
    add_heading3(doc, "2.2.1 주요 용어 정의")
    add_body(doc,
        "누적 비정상 수익률(CAR, Cumulative Abnormal Return): 이벤트 창(event window) 기간 동안 특정 기업의 "
        "실제 수익률에서 시장 모델로 예측한 정상 수익률을 차감한 값의 합. 공급망 리스크의 자본시장 파급 여부를 "
        "검증하는 핵심 지표이다."
    )
    add_body(doc,
        "공급망 리스크 이벤트(Supply Chain Risk Event): FinBERT 감성 분석 모델이 부정(NEG) 또는 긍정(POS) 감성 충격으로 "
        "분류한 뉴스 기사 클러스터. Z-score의 절댓값이 2.0을 초과하는 월을 이벤트 발생 시점으로 정의한다."
    )
    add_body(doc,
        "전파 관계(Contagion): 직접 충격을 받은 기업(DIRECT)으로부터 공급망 네트워크를 통해 연결된 기업들에게 "
        "CAR이 파급되는 현상. hop 수에 따라 1-hop(직접 연결), 2-hop(2단계 간접 연결) 전파로 구분한다."
    )
    add_body(doc,
        "GDELT GKG(Global Knowledge Graph): 전 세계 뉴스 미디어를 실시간 수집·분석하는 오픈소스 데이터베이스로, "
        "인물·조직·위치·테마 등의 개체 및 관계 정보를 구조화하여 제공한다."
    )

    add_heading3(doc, "2.2.2 개념적 분석 틀")
    add_body(doc,
        "본 연구의 개념적 틀은 세 층위로 구성된다. 첫 번째 층위는 뉴스 데이터에서 리스크 이벤트를 추출하는 "
        "데이터 파이프라인이다. 두 번째 층위는 이벤트 스터디를 통한 자본시장 파급 효과 측정이다. "
        "세 번째 층위는 지식 그래프와 GAT를 결합한 예측 및 전파 메커니즘 해석이다. "
        "이 세 층위의 통합이 본 연구의 핵심 기여이다."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 제3장 연구 설계 및 분석 방법
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "제3장  연구 설계 및 분석 방법")
    add_heading2(doc, "제1절  연구 대상 및 자료 수집")
    add_heading3(doc, "3.1.1 분석 기업 유니버스")
    add_body(doc,
        "본 연구의 분석 대상은 EV 배터리 공급망의 가치사슬 전반을 포괄하는 글로벌 상장 기업 44개사이다. "
        "기업 선정 기준은 (1) EV 배터리 가치사슬에서 핵심 사업 영역 보유, "
        "(2) 주요 거래소 상장으로 일별 주가 데이터 가용, (3) GDELT 기사에서 충분한 보도 건수 확보이다. "
        "이 중 GDELT 보도 커버리지와 주가 데이터 품질이 모두 충족되고 공급망 관계가 문헌으로 검증된 "
        "18개사(표 1의 * 표시)를 지식 그래프(KG) 및 Phase 5 GAT 분석의 핵심 기업으로 선정하였다. "
        "나머지 26개사는 Phase 4 CAR 이벤트 스터디의 전파(contagion) 분석에 포함된다."
    )
    add_table_from_data(doc,
        headers=["단계", "기업명", "티커", "거래소", "KG 핵심"],
        rows=[
            # Upstream/Refining
            ["Upstream", "Albemarle*",              "ALB",       "NYSE",   "●"],
            ["Upstream", "Glencore*",               "GLEN.L",    "LSE",    "●"],
            ["Upstream", "CMOC",                    "603993.SH", "SSE",    ""],
            ["Upstream", "Ganfeng Lithium",         "002460.SZ", "SZSE",   ""],
            ["Upstream", "Jiangxi Copper",          "600362.SH", "SSE",    ""],
            ["Upstream", "Tianqi Lithium",          "002466.SZ", "SZSE",   ""],
            ["Upstream", "SQM",                     "SQM",       "NYSE",   ""],
            # Materials
            ["Materials", "BASF*",                  "BAS.DE",    "XETRA",  "●"],
            ["Materials", "CNGR Advanced Material*","300919.SZ", "SZSE",   "●"],
            ["Materials", "EcoPro BM*",             "247540.KQ", "KOSDAQ", "●"],
            ["Materials", "Huayou Cobalt*",         "603799.SH", "SSE",    "●"],
            ["Materials", "LG Chem*",               "051910.KS", "KRX",    "●"],
            ["Materials", "Asahi Kasei",            "3407.T",    "TSE",    ""],
            ["Materials", "BTR New Material",       "920185.BJ", "BSE",    ""],
            ["Materials", "GEM",                    "002340.SZ", "SZSE",   ""],
            ["Materials", "Jiayuan Technology",     "688388.SH", "STAR",   ""],
            ["Materials", "L&F",                    "066970.KS", "KOSDAQ", ""],
            ["Materials", "Nuode",                  "600110.SH", "SSE",    ""],
            ["Materials", "POSCO Future M",         "003670.KS", "KRX",    ""],
            ["Materials", "Putailai",               "603659.SH", "SSE",    ""],
            # Cell/Pack
            ["Cell/Pack", "CALB*",                  "3931.HK",   "HKEX",   "●"],
            ["Cell/Pack", "CATL*",                  "300750.SZ", "SZSE",   "●"],
            ["Cell/Pack", "EVE Energy*",             "300014.SZ", "SZSE",   "●"],
            ["Cell/Pack", "Gotion High-Tech*",       "002074.SZ", "SZSE",   "●"],
            ["Cell/Pack", "LG Energy Solution*",     "373220.KS", "KRX",    "●"],
            ["Cell/Pack", "Capchem",                "300037.SZ", "SZSE",   ""],
            ["Cell/Pack", "Sunwoda",                "300207.SZ", "SZSE",   ""],
            # OEM
            ["OEM", "BMW*",                         "BMW.DE",    "XETRA",  "●"],
            ["OEM", "BYD*",                         "1211.HK",   "HKEX",   "●"],
            ["OEM", "Honda*",                       "7267.T",    "TSE",    "●"],
            ["OEM", "Mercedes-Benz*",               "MBG.DE",    "XETRA",  "●"],
            ["OEM", "Tesla*",                       "TSLA",      "NASDAQ", "●"],
            ["OEM", "Volkswagen*",                  "VOW3.DE",   "XETRA",  "●"],
            ["OEM", "BAIC BluePark",                "600733.SH", "SSE",    ""],
            ["OEM", "Changan Automobile",           "000625.SZ", "SZSE",   ""],
            ["OEM", "Dongfeng Motor",               "600006.SH", "SSE",    ""],
            ["OEM", "Ford",                         "F",         "NYSE",   ""],
            ["OEM", "GAC Group",                    "601238.SH", "SSE",    ""],
            ["OEM", "Geely Automobile",             "0175.HK",   "HKEX",   ""],
            ["OEM", "General Motors",               "GM",        "NYSE",   ""],
            ["OEM", "Great Wall Motor",             "601633.SH", "SSE",    ""],
            ["OEM", "JAC Motors",                   "600418.SH", "SSE",    ""],
            ["OEM", "Lucid",                        "LCID",      "NASDAQ", ""],
            ["OEM", "Rivian",                       "RIVN",      "NASDAQ", ""],
            ["OEM", "Toyota",                       "7203.T",    "TSE",    ""],
        ],
        caption="[표 1] EV 배터리 공급망 분석 기업 유니버스(44개사)\n* KG 핵심 기업(18개사, Phase 5 GAT 포함). 출처: 저자 작성"
    )

    add_heading3(doc, "3.1.2 GDELT GKG 뉴스 데이터")
    add_body(doc,
        "뉴스 데이터는 GDELT(Global Database of Events, Language, and Tone) GKG 데이터베이스에서 수집하였다. "
        "GDELT는 전 세계 100개국 이상의 미디어 기사를 15분 단위로 수집하는 오픈소스 뉴스 인텔리전스 플랫폼으로, "
        "Google BigQuery를 통해 SQL 방식으로 대용량 질의가 가능하다. "
        "수집 기간은 2018년 1월부터 2025년 12월까지이며, 분석 대상 기업명 및 공급망 관련 키워드(battery, lithium, "
        "cobalt, supply chain 등)가 포함된 기사를 BigQuery API로 추출하였다."
    )

    # 연도별 수집 현황 표
    add_table_from_data(doc,
        headers=["수집 연도", "URL 건수", "비고"],
        rows=[
            ["2018", "1,181,698", "수집 기반 구축 기간"],
            ["2019", "987,836",   ""],
            ["2020", "828,899",   "COVID-19 기간"],
            ["2021", "813,455",   ""],
            ["2022", "737,805",   "러-우 전쟁 발발"],
            ["2023", "899,617",   "미-중 갈등 심화"],
            ["2024", "869,683",   "희토류 수출 규제"],
            ["2025", "800,995",   "트럼프 관세 정책"],
            ["합계", "7,119,988", "고품질 5,855,113건(65.5%)"],
        ],
        caption="[표 2] GDELT GKG 연도별 수집 현황\n출처: 저자 수집 (GDELT BigQuery API)"
    )

    add_body(doc,
        "수집된 기사는 URL 품질(slug quality) 기준으로 고품질(good) 65.5%, 보통(fair) 22.6%, "
        "불명확(opaque) 11.6%, 공백(empty) 0.3%로 분류되었다. "
        "언어 분포는 영어 75.6%, 중국어 3.8%, 독일어 3.1%, 러시아어 2.0%, 우크라이나어 1.9%, "
        "기타 13.6%로, 다국어 뉴스 환경을 반영한다. "
        "실제 FinBERT 감성 분석에는 영어 기사 3,393,773건과 다국어(중·독·일 등) 기사 447,696건, "
        "합계 3,841,469건이 투입되었으며, 이 결과가 기업별 월별 감성 시계열 구축의 원천이 된다."
    )

    add_figure(doc, "C1_data_coverage.png",
               "[그림 1] GDELT 데이터 수집 커버리지 — 기업별·연도별 기사 건수\n출처: 저자 산출",
               desc_text=(
                   "각 셀은 해당 기업·연도의 GDELT 기사 수집 건수를 나타낸다. "
                   "CATL·BYD 등 중국 기업은 2022년 이후 급격히 보도량이 증가하였으며, "
                   "Tesla·BMW 등 OEM 기업은 전 기간에 걸쳐 안정적인 보도가 유지되었다. "
                   "CNGR(중국 소재사) 등 일부 기업은 보도량이 극히 적어 분석 통계력에 제약이 존재한다."
               ))

    add_heading3(doc, "3.1.3 주가 및 벤치마크 데이터")
    add_body(doc,
        "주가 데이터는 Yahoo Finance API를 통해 수집한 일별 수정 종가(adjusted close price)를 사용하였다. "
        "수집 기간은 2017년 1월부터 2026년 5월까지이며, 총 92,350건의 수익률 관측치를 확보하였다. "
        "결측률은 38개 기업 전체에서 0%로 양호하며, 수익률 절댓값 20% 초과 극단값 비율도 0.1% 미만이다."
    )
    add_body(doc,
        "각 기업의 거래소별 시장 벤치마크 인덱스를 정상 수익률 추정에 활용하였다. "
        "한국 기업: KOSPI, 중국 A주: CSI300, 중국 H주(홍콩): HSI, 미국: S&P500, 독일: DAX30, 일본: Nikkei225, "
        "영국(LSE): FTSE100. 이는 각 거래소의 거래일 스케줄 차이를 통제하기 위한 설계이다."
    )

    add_heading3(doc, "3.1.4 데이터 전처리 및 품질 검증")
    add_body(doc,
        "GDELT 원본 15,906,417건의 URL 중 기업 관련 슬러그 매칭 및 중복 제거를 거쳐 "
        "7,119,988건의 고유 URL을 확보하였다. 이 중 고품질(good quality) URL 5,855,113건이 FinBERT 처리 후보로 분류되었으나, "
        "실제 FinBERT 감성 분석에 투입된 건수는 영어 3,393,773건 + 다국어 447,696건 = 총 3,841,469건이다. "
        "나머지는 기업명 매칭 불일치 또는 처리 우선순위 절충으로 제외되었다. "
        "기업별 언급 매칭은 Aho-Corasick 알고리즘 기반 멀티패턴 검색으로 수행하였으며, "
        "기업명 동의어(약어, 현지어 표기 포함)를 시드 사전으로 구축하여 정밀도를 높였다."
    )
    add_body(doc,
        "최종적으로 FinBERT 처리 후 18개 기업 × 101개월의 감성 시계열(tone_monthly_zscore)이 생성되었다. "
        "이 중 Z-score 절댓값 ≥ 2.0을 충족하는 음성 충격 37건, 양성 충격 28건이 CAR 이벤트 스터디의 입력 이벤트로 확정되었다."
    )

    add_heading2(doc, "제2절  분석 도구 및 절차")
    add_heading3(doc, "3.2.1 FinBERT 감성 분석 및 이벤트 식별")
    add_body(doc,
        "수집된 뉴스 기사의 감성 분석에는 FinBERT(Araci, 2019)를 사용하였다. FinBERT는 BERT 사전학습 언어 모델을 "
        "금융 도메인 텍스트로 파인튜닝한 모델로, 금융 뉴스에서 긍정/부정/중립 감성을 높은 정확도로 분류한다. "
        "중문·독문 기사에는 multilingual-FinBERT를 적용하여 언어 커버리지를 확장하였다."
    )
    add_body(doc,
        "이벤트 식별 절차: "
        "① 기업별 월별 감성 점수의 기사 수 가중 평균 산출, "
        "② 36개월 롤링 평균·표준편차 기반 Z-score 정규화 (분포 정상화 목적), "
        "③ |Z-score| ≥ 2.0인 월을 감성 충격 이벤트로 정의 — 통계적으로 전체 분포의 상위 2.3%에 해당하는 이상 이벤트."
    )

    add_heading3(doc, "3.2.2 CAR 이벤트 스터디 및 베이스라인")
    add_body(doc,
        "이벤트 스터디는 Brown and Warner(1985)의 시장 모델(market model)을 베이스라인으로 수행하였다. "
        "정상 수익률 추정 창(estimation window)은 이벤트 발생 전 [-120, -21] 영업일이며, "
        "이 구간의 OLS 회귀(Ri = αi + βi·Rm + εi)로 기대 수익률 파라미터(α, β)를 추정한다. "
        "이벤트 창([0,+21]) 기간의 실제 수익률에서 추정 정상 수익률을 차감한 값이 비정상 수익률(AR)이며, "
        "이를 누적한 값이 CAR이다. 즉, 시장 전체 움직임을 제거한 '순수 기업별 충격' 크기를 측정하는 구조이다."
    )
    add_body(doc,
        "통계 검정은 횡단면 t-검정과 부호 검정을 병행하였으며, 이벤트-월 클러스터 부트스트랩(B=1,000)으로 "
        "표준 오차의 견고성을 확보하였다. 공급망 전파 효과는 직접 충격 기업(DIRECT)으로부터 "
        "hop 거리별(1-hop, 2-hop) 연결 기업의 CAR을 유형별로 분류하여 비교하였다."
    )

    add_heading3(doc, "3.2.3 공급망 지식 그래프 구축")
    add_body(doc,
        "공급망 지식 그래프는 학술 문헌, 기업 연간보고서, 공급망 데이터베이스를 수동 검토하여 구축하였다. "
        "그래프는 두 층위로 구성된다."
    )
    add_body(doc,
        "【전체 KG】 분석 유니버스 44개 기업 노드에 국가(11개)·핵심 광물(6종)·주요 항구(2개) 등 "
        "비기업 노드 25개를 추가한 총 69개 노드, 287개 엣지로 구성된다. "
        "엣지 유형은 SUPPLIES(164), LOCATED_IN(54), PARTNERS_WITH(29), SOURCES_FROM(10), "
        "COMPETES_WITH(10), BUYS_FROM(9), PRODUCES(7), OWNS(2), DEPENDS_ON_ROUTE(2)이다. "
        "이 전체 KG는 공급망 리스크 경로와 hop 거리 계산의 기반 구조로 활용된다."
    )
    add_body(doc,
        "【GAT 서브그래프】 Phase 5 Temporal GAT 학습에는 18개 핵심 기업 간 내부 엣지만 추출한 "
        "서브그래프(18 노드, 65 엣지)를 사용하였다. "
        "엣지 구성: SUPPLIES 50건, PARTNERS_WITH 7건, COMPETES_WITH 5건, BUYS_FROM 2건, OWNS 1건. "
        "서브그래프로 한정한 이유는 GAT 입력 피처(18개 기업 감성 시계열)와 노드 집합을 일치시키기 위함이다."
    )
    add_figure(doc, "fig_kg_subgraph_final.png",
               "[그림 2] EV 배터리 공급망 지식 그래프 — GAT 서브그래프 (18개 노드, 65개 엣지)\n출처: 저자 작성",
               desc_text=(
                   "노드 색상은 공급망 단계(Upstream=갈색, Materials=주황, Cell/Pack=파랑, OEM=초록)를 나타낸다. "
                   "엣지 방향은 SUPPLIES(공급) 관계의 흐름이며, 양방향 엣지는 PARTNERS_WITH 또는 COMPETES_WITH 관계이다. "
                   "CATL·BYD 등 중국 기업이 허브(hub) 역할을, Glencore·Albemarle 등 Upstream 기업이 "
                   "다수 Cell/Pack 기업과 연결되어 병목(bottleneck) 구조를 형성한다. "
                   "전체 KG(44개 기업 + 25개 비기업 노드, 287 엣지)의 핵심 서브셋이다."
               ))

    add_heading3(doc, "3.2.4 그래프 어텐션 네트워크(GAT) 및 베이스라인 설계")
    add_body(doc,
        "공급망 구조 정보의 예측 기여도 검증을 위해 Veličković et al.(2018)의 GAT 모델을 적용하였다. "
        "아키텍처는 GRU(hidden=32) 시계열 인코더와 GAT(hidden=16, head=4) 그래프 레이어를 결합한 구조이며, "
        "입력 피처는 FinBERT 감성 Z-score를 포함한 18개 변수이다. "
        "예측 목표는 다음 월의 비정상 수익률 방향성(상승/하락 이진 분류)이다."
    )
    add_body(doc,
        "베이스라인 모델 구성 — Ablation Study의 세 모델은 다음과 같다. "
        "① Real Graph GAT: 실제 공급망 구조 65개 엣지를 인접 행렬로 사용 (제안 모델). "
        "② Random Graph GAT: 동일 밀도의 무작위 그래프로 대체 — 그래프 연결 구조 자체의 기여도 검증. "
        "③ No Graph GRU: 그래프 레이어를 완전히 제거한 순수 시계열 베이스라인 — "
        "공급망 정보 없이 뉴스 감성만으로 예측 가능한 상한선(ceiling)을 정의한다."
    )
    add_body(doc,
        "과적합 통제를 위해 Walk-Forward Cross-Validation(WFCV)을 적용하였다. "
        "훈련 창(Train): 2017-01~2023-06 (68개월), 검증 창(Val): 2023-07~2024-06 (12개월), "
        "테스트 창(Test): 2024-07~2026-05 (23개월). "
        "테스트 기간은 트럼프 관세 정책(2025년), 미-중 무역 갈등 심화 등 실제 리스크 이벤트가 집중된 시기로, "
        "모델의 실전 예측력 검증에 적합한 구간이다."
    )
    add_figure(doc, "fig7_wfcv.png",
               "[그림 3] Walk-Forward Cross-Validation(WFCV) 설계\n출처: 저자 작성",
               desc_text=(
                   "각 슬라이딩 윈도우는 훈련 데이터를 확장하며 테스트 기간을 1개월씩 전진하는 방식이다. "
                   "이를 통해 미래 데이터 누수(data leakage) 없이 모델의 시간 외 예측 성능을 평가한다. "
                   "총 23개 테스트 월(2024-07~2026-05)에 걸친 평균 IC와 방향 정확도를 최종 성능 지표로 사용하였다."
               ))

    # ══════════════════════════════════════════════════════════════════════════
    # 제4장 연구 결과 및 실증 분석
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "제4장  연구 결과 및 실증 분석")
    add_heading2(doc, "제1절  자료의 특성")
    add_heading3(doc, "4.1.1 뉴스 이벤트 단계별 집계")
    add_table_from_data(doc,
        headers=["처리 단계", "건수", "비고"],
        rows=[
            ["GDELT 원본 수집 URL",           "15,906,417건", "전체 GKG 레코드"],
            ["고유 URL (중복 제거)",           "7,119,988건",  ""],
            ["고품질 URL 후보 (good quality)", "5,855,113건",  "전체의 65.5% — FinBERT 입력 후보"],
            ["실제 FinBERT 처리 (영어)",       "3,393,773건",  "finbert_v2.parquet"],
            ["실제 FinBERT 처리 (다국어)",     "447,696건",    "finbert_v2_multilingual.parquet"],
            ["FinBERT 처리 합계",              "3,841,469건",  "기업명 매칭 통과 분"],
            ["기업별 월별 이벤트 집계",        "40,928건",     "Aho-Corasick + Z-score 집계"],
            ["FinBERT 음성 충격(Z≤-2.0)",     "37건",         "NEG 이벤트 (Phase 4 입력)"],
            ["FinBERT 양성 충격(Z≥+2.0)",     "28건",         "POS 이벤트 (Phase 4 입력)"],
        ],
        caption="[표 2] GDELT 뉴스 데이터 단계별 집계 요약\n출처: 저자 작성 (GDELT GKG, FinBERT 분석 결과)"
    )

    add_heading3(doc, "4.1.2 주가 데이터 특성")
    add_body(doc,
        "44개 기업의 일별 주가 데이터는 총 92,350건의 수익률 관측치를 포함한다. "
        "결측률은 38개 기업 전체에서 0%로 양호하며, 수익률 절댓값 20% 초과 극단값 비율은 0.1% 미만이다. "
        "예외적으로 LCID(Lucid Group)가 0.9%의 극단값 비율로 가장 높았으며, "
        "이는 소형 신규 EV 기업 특유의 높은 변동성을 반영한다. 분석 기간 내 전체 수익률 분포는 정규성 검정을 통과하여 "
        "시장 모델 기반 이벤트 스터디 적용에 적합한 수준임을 확인하였다."
    )

    add_heading2(doc, "제2절  분석 결과")
    add_heading3(doc, "4.2.1 직접 효과(Direct Effect) — CAR 이벤트 스터디")
    add_body(doc,
        "시장 모델 베이스라인 대비 비정상 수익률(AR) 분석 결과, 주 분석 창 [0,+21]에서 "
        "전체 이벤트 기준 CAAR=+0.62%(p=0.544)로 통계적으로 유의하지 않았다. "
        "이는 EV 배터리 뉴스 감성 충격이 직접 충격 기업의 주가에 즉각적이고 체계적인 영향을 미치지 않음을 보여준다."
    )
    add_table_from_data(doc,
        headers=["이벤트 창", "N", "CAAR", "t-통계량", "p값", "해석"],
        rows=[
            ["[-1, +1]", "129", "+0.29%", "0.637", "0.525", "비유의"],
            ["[0, +1]",  "129", "+0.11%", "0.269", "0.789", "비유의"],
            ["[0, +5]",  "129", "+0.43%", "0.567", "0.572", "비유의"],
            ["[0, +10]", "129", "+0.52%", "0.522", "0.603", "비유의"],
            ["[0, +21]", "129", "+0.62%", "0.608", "0.544", "비유의"],
        ],
        caption="[표 3] 이벤트 창별 CAAR 추정 결과 (DIRECT, N=129)\n출처: 저자 산출 (Phase 4 CAR Event Study, 시장 모델 추정 [-120,-21])"
    )
    add_body(doc,
        "충격 유형별 이질성(주 창 [0,+21], N=125, 44개사 기준): "
        "NEG shock(N=70): CAAR=-1.73%, t=-1.397, p=0.167 (비유의); "
        "POS shock(N=55): CAAR=+3.62%, t=+2.227, p=0.030** (유의). "
        "긍정 충격에 대한 직접 반응만 유의하였으며, 이는 EV 배터리 산업의 성장 기대감이 "
        "호재 정보를 시장에 빠르게 반영시키는 반면, 부정 충격에 대한 반응은 분산이 크고 불확실함을 시사한다."
    )
    add_figure(doc, "A3_neg_pos_comparison.png",
               "[그림 4] NEG vs POS 충격 유형별 CAAR 비교 (이벤트 창 [0,+21], 44개사)\n출처: 저자 산출",
               desc_text=(
                   "막대는 각 충격 유형(NEG, POS)의 CAAR을, 오차막대는 95% 신뢰구간을 나타낸다. "
                   "POS 충격(CAAR=+3.62%, p=0.030**)만 통계적으로 유의하며, "
                   "NEG 충격은 CAAR=-1.73%으로 음의 방향이지만 통계적으로 유의하지 않다(p=0.167). "
                   "이는 시장이 긍정 정보를 부정 정보보다 더 즉각적으로 반영하는 정보 비대칭 패턴을 보여준다."
               ))

    add_heading3(doc, "4.2.2 공급망 전파 효과(Contagion Effect)")
    add_body(doc,
        "직접 충격 기업(DIRECT)의 공급망 연결 기업들에게 CAR이 전파되는지를 hop 거리별로 분석하였다. "
        "1-hop 연결 기업(DOWNSTREAM, UPSTREAM, PEER)에서는 유의미한 전파 효과가 관찰되지 않았으나, "
        "2-hop 간접 연결 기업에서 강건한 통계적 유의성이 확인되었다."
    )
    add_table_from_data(doc,
        headers=["전파 유형", "hop", "N", "CAAR", "t-통계량", "p값", "유의성"],
        rows=[
            ["DIRECT",      "-", "125",  "+0.62%", "+0.608", "0.544",  "-"],
            ["DOWNSTREAM",  "1", "161",  "-0.43%", "-0.442", "0.659",  "-"],
            ["UPSTREAM",    "1", "203",  "+0.05%", "+0.060", "0.952",  "-"],
            ["PEER",        "1", "96",   "+0.58%", "+0.391", "0.697",  "-"],
            ["MIXED_2",     "2", "3,356","+0.90%", "+3.698", "<0.001", "***"],
            ["UPSTREAM_2",  "2", "898",  "+1.17%", "+2.300", "0.022",  "**"],
            ["DOWNSTREAM_2","2", "384",  "-0.49%", "-0.764", "0.445",  "-"],
            ["PEER_2",      "2", "44",   "+5.11%", "+1.803", "0.078",  "*"],
        ],
        caption="[표 4] 공급망 hop 거리별 전파 효과 (이벤트 창 [0, +21], 44개사)\n출처: 저자 산출 (*** p<0.001, ** p<0.05, * p<0.10)"
    )
    add_body(doc,
        "1-hop 직접 연결 기업(DOWNSTREAM, UPSTREAM, PEER)에서는 모두 통계적으로 유의한 전파 효과가 없었다. "
        "반면, 2-hop 간접 연결 기업에서는 MIXED_2(N=3,356, CAAR=+0.90%, p<0.001)와 "
        "UPSTREAM_2(N=898, CAAR=+1.17%, p=0.022)의 강건한 유의성이 확인되었다. "
        "PEER_2(N=44, CAAR=+5.11%, p=0.078)는 약한 유의성이 관찰되었으나 표본이 소규모이다. "
        "DOWNSTREAM_2(CAAR=-0.49%, p=0.445)는 비유의하였다."
    )
    add_figure(doc, "fig_contagion_final.png",
               "[그림 5] 공급망 hop별 전파 효과 — hop 거리·유형별 CAAR 비교 (이벤트 창 [0,+21], 44개사)\n출처: 저자 산출",
               desc_text=(
                   "x축은 공급망 hop 거리 및 관계 유형, y축은 CAAR(%)이다. "
                   "1-hop 기업들은 모두 통계적 비유의 구간에 머물러 있는 반면, "
                   "MIXED_2와 UPSTREAM_2는 95% 신뢰구간이 0을 상회하여 명확한 유의성을 보인다. "
                   "이 패턴은 '정보 마찰(information friction)' 가설과 일치하며, "
                   "시장이 1차 공급망 파급을 즉각 소화하지 못한 채 2단계에서 더 강하게 반응하는 구조를 드러낸다."
               ))

    add_heading3(doc, "4.2.3 강건성 검정(Robustness Check)")
    add_body(doc,
        "핵심 결과의 강건성 검증을 위해 세 가지 대안 설계를 적용하였다. "
        "첫째, T0 대안(이벤트 발생월 1일 T0 설정): NEG shock에서 CAAR=-3.18%(p=0.070*)로 방향이 역전되며 "
        "약한 유의성이 확인되었다. "
        "둘째, 클러스터 표준오차(이벤트-월 클러스터, B=1,000) 적용 후 p값이 전반적으로 상승하였으나, "
        "MIXED_2와 UPSTREAM_2의 유의성은 유지되어 핵심 결과의 강건성이 지지된다. "
        "셋째, 코로나19 기간(2020~2021) 제외 분석에서도 2-hop 전파 패턴의 방향성이 유지되었다."
    )

    add_heading3(doc, "4.2.4 GAT vs 베이스라인 성능 비교")
    add_body(doc,
        "Ablation Study를 통해 세 모델의 성능을 직접 비교하였다. "
        "테스트 기간(2024-07~2026-05, 23개월, N=18 기업) 기준 평가 지표는 "
        "AUC(방향 예측 정확도), F1(균형 정확도), IC(Information Coefficient, 예측-실현 상관계수)이다."
    )
    add_table_from_data(doc,
        headers=["모델", "AUC", "F1", "Acc", "IC 월평균", "IC>0 비율", "해석"],
        rows=[
            ["Real Graph GAT\n(제안 모델)", "0.5281", "0.5667", "48.0%", "+0.166", "67%",
             "실제 공급망 구조 활용"],
            ["Random Graph GAT\n(구조 기여 검증)", "0.5044", "0.5778", "46.8%", "-0.104", "20%",
             "무작위 엣지 연결"],
            ["No Graph GRU\n(순수 시계열 베이스라인)", "0.4594", "0.4825", "46.8%", "-0.002", "60%",
             "그래프 미사용"],
        ],
        caption="[표 5] GAT Ablation Study — 베이스라인 대비 성능 비교 (테스트 2024-07~2026-05)\n출처: 저자 산출 (Permutation test IC p=0.398)"
    )
    add_body(doc,
        "Real Graph GAT(AUC=0.5281)는 Random Graph GAT(0.5044)보다 AUC 기준 +2.4%p 우월하며, "
        "GRU 베이스라인(0.4594) 대비 +6.9%p 높은 수준이다. "
        "이는 실제 공급망 연결 구조(어떤 기업이 어떤 기업에 납품하는지)가 순수 시계열 피처보다 추가적인 예측 정보를 "
        "제공함을 보여주며, 무작위 그래프 대비 우위는 네트워크 위상 구조 자체의 기여임을 시사한다."
    )
    add_body(doc,
        "IC 지표에서 Real Graph GAT는 23개월 중 67%(15.4개월)에서 양의 IC를 기록한 반면, "
        "Random Graph GAT는 20%(4.6개월)에 불과하여, 실제 그래프 구조가 IC 안정성에 크게 기여함을 확인하였다. "
        "다만, Permutation test IC p=0.398로 통계적 유의성은 확인되지 않아, "
        "18개 기업 소규모 유니버스에서의 통계적 검정력 한계가 있다."
    )
    add_figure(doc, "fig_ablation_final.png",
               "[그림 6] GAT Ablation Study — 3개 모델 AUC·IC>0% 비교 (SEED=42, 18개사)\n출처: 저자 산출",
               desc_text=(
                   "왼쪽 패널은 세 모델의 AUC 분포(박스플롯)를 비교하며, 오른쪽 패널은 월별 IC 추이를 나타낸다. "
                   "Real Graph GAT(파랑)는 대부분의 테스트 월에서 Random GAT(주황)와 GRU(회색)보다 높은 IC를 유지한다. "
                   "Random Graph GAT의 IC가 음수(-0.104)로 떨어지는 구간은 무작위 엣지가 오히려 예측을 교란함을 의미하며, "
                   "이는 그래프 구조의 선택이 모델 성능에 결정적 영향을 미침을 반증한다."
               ))
    add_figure(doc, "fig_monthly_ic_final.png",
               "[그림 7] Real Graph GAT 월별 IC 추이 (Test: 2024-07~2026-05)\n출처: 저자 산출",
               desc_text=(
                   "양의 IC(파랑)는 해당 월에 GAT 예측 방향이 실현 수익률 방향과 일치한 것을 의미한다. "
                   "23개월 중 15개월(67%)에서 양의 IC가 관찰되었으며, 평균 IC=+0.166이다. "
                   "2025-04(IC=+0.591), 2024-11(IC=+0.507), 2025-09(IC=+0.515) 등 리스크 이벤트 집중 시기에 예측력이 높았다."
               ))
    add_figure(doc, "fig8_attention.png",
               "[그림 8] GAT 어텐션 가중치 분포 — 상위 공급망 엣지\n출처: 저자 산출",
               desc_text=(
                   "어텐션 가중치는 GAT 레이어가 예측 시 각 공급망 관계에 부여하는 상대적 중요도이다. "
                   "CNGR↔Albemarle(0.149), BYD↔Honda(0.148), EcoPro BM↔Huayou Cobalt(0.141) 순으로 높으며, "
                   "이는 소재(Materials) 기업과 OEM 간의 교차 공급 관계가 리스크 전파의 핵심 경로임을 시사한다. "
                   "BYD↔CALB, CATL↔Glencore 등 중국 내 수직계열화 엣지도 높은 어텐션을 받아 "
                   "중국 공급망 내부 전파의 중요성을 확인할 수 있다."
               ))

    add_heading2(doc, "제3절  결과에 대한 논의")
    add_body(doc,
        "첫째, 1-hop 직접 효과 부재 및 2-hop 간접 전파의 강건성은 시장의 정보 마찰(information friction) 가능성을 시사한다. "
        "EV 배터리 공급망의 복잡한 연결 구조에서, 직접 충격 기업의 1차 파트너보다 2차 파트너에게 충격이 더 명확하게 "
        "자본시장에 반영된다(MIXED_2 p<0.001, UPSTREAM_2 p=0.022). "
        "이는 시장 참여자들이 1-hop 연결 기업에 대한 정보를 즉각 처리하지 못하거나, "
        "2-hop 기업에서의 실물 경제 충격이 더 지연되어 나타나기 때문으로 해석할 수 있다."
    )
    add_body(doc,
        "둘째, POS 충격의 직접 효과(CAAR=+3.62%, p=0.030)만 유의하고 NEG 충격(-1.73%, p=0.167)은 "
        "비유의한 비대칭 패턴은, EV 배터리 산업의 성장 기대감이 긍정 정보의 신속한 가격 반영을 촉진하는 반면 "
        "부정 충격에 대한 시장 반응이 불확실하고 분산이 큼을 시사한다. "
        "월별 Z-score 기반 이벤트 정의는 이미 시장에 반영된 정보도 포함할 수 있어, "
        "T0 설계의 날짜 모호성이 NEG 반응의 희석에 기여했을 수 있다."
    )
    add_body(doc,
        "셋째, 그래프 구조의 탐색적 기여도는 공급망 네트워크 정보가 개별 기업의 시계열 정보를 보완하는 예측 신호를 "
        "제공함을 시사한다. 그러나 Permutation test p=0.398로 통계적 유의성은 확인되지 않아, "
        "소규모(18개 기업) 네트워크에서 GNN의 효과를 통계적으로 입증하기 위해서는 더 큰 유니버스가 필요하다."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 제5장 결론
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "제5장  결  론")
    add_heading2(doc, "제1절  연구 요약")
    add_body(doc,
        "본 연구는 EV 배터리 공급망 리스크의 자본시장 파급 메커니즘을 실증적으로 규명하기 위해 "
        "GDELT 뉴스 기반 리스크 이벤트 식별, CAR 이벤트 스터디, 그래프 어텐션 네트워크를 활용한 "
        "End-to-End 분석 파이프라인을 구축하였다."
    )
    add_body(doc, "핵심 결과는 다음과 같다.")
    add_body(doc,
        "① 전체 이벤트 기준 직접 효과 CAAR=+0.62%(p=0.544)로 비유의하였으나, "
        "충격 유형별로는 POS shock CAAR=+3.62%(p=0.030**)만 유의하여 시장의 비대칭적 반응이 확인되었다."
    )
    add_body(doc,
        "② 2-hop 간접 전파에서 MIXED_2(CAAR=+0.90%, p<0.001)와 UPSTREAM_2(CAAR=+1.17%, p=0.022)의 "
        "강건한 유의성이 확인되어, 정보 마찰에 의한 지연 전파 가설을 지지한다."
    )
    add_body(doc,
        "③ Ablation Study를 통해 실제 공급망 그래프 구조(AUC=0.5281)가 랜덤 그래프(0.5044) 및 "
        "그래프 미사용 모델(0.4594)보다 우월한 예측 성능을 보였으나, 소규모 유니버스로 인해 "
        "통계적 유의성 확보에는 한계가 있었다."
    )

    add_heading2(doc, "제2절  정책적 제언 및 시사점")
    add_body(doc,
        "투자자 관점: EV 배터리 공급망 리스크는 직접 충격 기업뿐 아니라 2-hop 간접 연결 기업의 포트폴리오에도 "
        "유의미한 영향을 미친다. 투자자들은 EV 관련 기업 분석 시 직접 거래 파트너를 넘어 2단계 이상의 공급망 "
        "노출도를 고려한 리스크 관리 전략을 수립할 필요가 있다."
    )
    add_body(doc,
        "기업 관점: 공급망 가시성(supply chain visibility) 확보가 리스크 관리의 핵심 과제임을 시사한다. "
        "UPSTREAM_2 기업군에서 양(+)의 비정상 수익률이 관찰된 결과는, 직접 충격 기업의 대안 공급처로 "
        "부상하는 효과를 반영하며, 공급망 다변화 전략의 실질적 가치를 보여준다."
    )
    add_body(doc,
        "정책 관점: EV 배터리 핵심 광물의 특정 국가·기업 집중도가 높은 현황에서, "
        "공급망 다변화 정책은 시스템적 리스크 전파를 억제하는 효과를 지닐 수 있다. "
        "본 연구의 GDELT 기반 실시간 뉴스 감성 파이프라인은 정책 당국의 공급망 조기경보(EWS) 시스템 구축에도 "
        "활용 가능하다."
    )

    add_heading2(doc, "제3절  연구의 한계 및 향후 과제")
    add_body(doc, "연구의 한계는 다음과 같다.")
    add_body(doc, "① 44개 기업·129건 이벤트의 표본 규모 한계 — 클러스터 SE 적용 후 유의성 약화.")
    add_body(doc, "② 월별 감성 충격 기반 이벤트 정의의 날짜 특정 모호성 — 일별 고빈도 분석으로 개선 가능.")
    add_body(doc, "③ 18개 기업 소규모 네트워크에서 GNN의 통계적 유의성 확보 한계.")
    add_body(doc, "④ 상관관계 기반 분석으로 인과 관계 규명의 한계 — DID 설계 등으로 보완 필요.")
    add_body(doc, "향후 과제는 다음과 같다.")
    add_body(doc, "① 일별 뉴스-주가 고빈도 분석을 통한 이벤트 식별력 강화,")
    add_body(doc, "② 100개 이상 기업으로의 공급망 유니버스 확장 및 GDELT 커버리지 검증,")
    add_body(doc, "③ Temporal GNN(TGAT, DyRep) 적용을 통한 동적 엣지 가중치 모델링,")
    add_body(doc, "④ 반도체·희토류 등 타 산업으로의 방법론 일반화.")

    # ══════════════════════════════════════════════════════════════════════════
    # 참고문헌
    # ══════════════════════════════════════════════════════════════════════════
    add_heading1(doc, "참고문헌")
    refs = [
        "Acemoglu, D., Carvalho, V. M., Ozdaglar, A., & Tahbaz-Salehi, A. (2012). The network origins of aggregate fluctuations. Econometrica, 80(5), 1977–2016.",
        "Araci, D. (2019). FinBERT: Financial sentiment analysis with pre-trained language models. arXiv preprint arXiv:1908.10063.",
        "Baker, S. R., Bloom, N., & Davis, S. J. (2016). Measuring economic policy uncertainty. The Quarterly Journal of Economics, 131(4), 1593–1636.",
        "Bakshi, N., & Kleindorfer, P. (2009). Co-opetition and investment for supply-chain resilience. Production and Operations Management, 18(6), 583–603.",
        "Brown, S. J., & Warner, J. B. (1985). Using daily stock returns: The case of event studies. Journal of Financial Economics, 14(1), 3–31.",
        "Carvalho, V. M. (2014). From micro to macro via production networks. Journal of Economic Perspectives, 28(4), 23–48.",
        "Gulley, A. L., Nassar, N. T., & Phillip, S. (2018). China, the United States, and competition for resources that enable emerging technologies. Proceedings of the National Academy of Sciences, 115(16), 4111–4115.",
        "Hendricks, K. B., & Singhal, V. R. (2003). The effect of supply chain glitches on shareholder wealth. Journal of Operations Management, 21(5), 501–522.",
        "Jüttner, U. (2005). Supply chain risk management: Understanding the business requirements from a practitioner perspective. The International Journal of Logistics Management, 16(1), 120–141.",
        "Kipf, T. N., & Welling, M. (2017). Semi-supervised classification with graph convolutional networks. ICLR 2017.",
        "Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global data on events, location, and tone, 1979–2012. ISA Annual Convention.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. The Journal of Finance, 62(3), 1139–1168.",
        "Veličković, P., Cucurull, G., Casanova, A., Romero, A., Lio, P., & Bengio, Y. (2018). Graph attention networks. ICLR 2018.",
    ]
    for ref in refs:
        p = add_para(doc, ref, F_LIGHT, 10, space_before=0, space_after=3, line_spacing=Pt(16))
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)


if __name__ == "__main__":
    build()
