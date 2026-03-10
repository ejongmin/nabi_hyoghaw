from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional
import pandas as pd

from src.tools.seed_from_xlsx import seed_edges_from_excel, SeedImportResult

def seed_edges_from_docx_tables(docx_path: Path,
                                temp_xlsx_path: Path,
                                *args, **kwargs) -> SeedImportResult:
    """Best-effort DOCX table import.

    Strategy:
    - Parse all tables from DOCX
    - Export as an XLSX (one sheet per table)
    - Reuse the Excel import pipeline (seed_edges_from_excel) per table
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"docx not found: {docx_path}")

    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx 미설치: pip install python-docx") from e

    doc = Document(str(docx_path))
    if not doc.tables:
        raise ValueError("DOCX에 표(table)가 없습니다. 표 형태로 정리된 문서만 자동 변환 가능합니다.")

    # Build sheets
    with pd.ExcelWriter(temp_xlsx_path, engine="openpyxl") as writer:
        for i, table in enumerate(doc.tables):
            rows = []
            for r in table.rows:
                rows.append([c.text.strip() for c in r.cells])
            if not rows:
                continue
            header = rows[0]
            body = rows[1:] if len(rows) > 1 else []
            df = pd.DataFrame(body, columns=header)
            df.to_excel(writer, sheet_name=f"table_{i+1}", index=False)

    # Use first sheet by default; user can rerun specifying sheet
    if "sheet" not in kwargs:
        kwargs["sheet"] = 0
    return seed_edges_from_excel(temp_xlsx_path, **kwargs)
